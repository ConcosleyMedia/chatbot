import os
import pickle
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document as DocxDocument

RAW_DIR = os.path.join(os.path.dirname(__file__), '../raw')
INDEX_DIR = os.path.join(os.path.dirname(__file__), '../faiss_index/')
CHUNK_SIZE = 2000  # ~500 tokens

load_dotenv()

def load_docs():
    docs = []
    # Recursively walk through RAW_DIR and all subdirectories
    for root, _, files in os.walk(RAW_DIR):
        for fname in files:
            path = os.path.join(root, fname)
            if fname.endswith('.md'):
                with open(path, 'r', encoding='utf-8') as f:
                    docs.append(Document(page_content=f.read(), metadata={"source": os.path.relpath(path, RAW_DIR)}))
            elif fname.endswith('.pdf'):
                reader = PdfReader(path)
                text = "\n".join(page.extract_text() or '' for page in reader.pages)
                docs.append(Document(page_content=text, metadata={"source": os.path.relpath(path, RAW_DIR)}))
            elif fname.endswith('.html'):
                with open(path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'html.parser')
                    text = soup.get_text()
                    docs.append(Document(page_content=text, metadata={"source": os.path.relpath(path, RAW_DIR)}))
            elif fname.endswith('.docx'):
                doc = DocxDocument(path)
                text = "\n".join([para.text for para in doc.paragraphs])
                docs.append(Document(page_content=text, metadata={"source": os.path.relpath(path, RAW_DIR)}))
    return docs

def main():
    docs = load_docs()
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=200)
    all_chunks = []
    for doc in docs:
        for chunk in splitter.split_text(doc.page_content):
            all_chunks.append(Document(page_content=chunk, metadata=doc.metadata))
    print(f"Loaded {len(all_chunks)} chunks.")
    embeddings = OpenAIEmbeddings()
    # Batch the chunks to avoid exceeding token limits
    batch_size = 100
    vectordb = None
    for i in range(0, len(all_chunks), batch_size):
        batch = all_chunks[i:i+batch_size]
        if not batch:
            continue
        if vectordb is None:
            vectordb = FAISS.from_documents(batch, embeddings)
        else:
            vectordb.add_documents(batch)
        print(f"Processed batch {i // batch_size + 1} ({len(batch)} docs)")
    if vectordb:
        vectordb.save_local(INDEX_DIR)
        print(f"FAISS index saved to {INDEX_DIR}")
    else:
        print("No documents were indexed!")

if __name__ == "__main__":
    main()
